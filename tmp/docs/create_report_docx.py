from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path("/Users/trunovatatyana/Documents/test/climo")
SOURCE = ROOT / "output" / "technical_implementation_report.md"
TARGET = ROOT / "output" / "doc" / "technical_implementation_report.docx"


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Menlo"
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.space_after = Pt(8)


def main():
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    for name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True

    section = document.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(56)
    section.right_margin = Pt(56)

    in_code = False
    code_lines = []

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line.startswith("# "):
            paragraph = document.add_heading(line[2:], level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if line.startswith("## "):
            document.add_heading(line[3:], level=2)
            continue

        if line.startswith("### "):
            document.add_heading(line[4:], level=3)
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
            continue

        numbered = line.split(". ", 1)
        if len(numbered) == 2 and numbered[0].isdigit():
            document.add_paragraph(numbered[1], style="List Number")
            continue

        document.add_paragraph(line)

    if code_lines:
        add_code_block(document, code_lines)

    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
